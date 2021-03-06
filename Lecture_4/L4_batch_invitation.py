import locale
import os
import time
from os import makedirs
from os.path import exists

import docx
import pandas as pd
import xlwings as xw
from docx.enum.text import WD_ALIGN_PARAGRAPH  # 设置对齐方式
from docx.oxml.ns import qn  # 设置字体
from docx.shared import Cm, Inches, Pt, RGBColor  # 用来设置颜色、字体大小、缩进


def read_info_xlwings(filename):
    app = xw.App(visible=True, add_book=False)
    wb = app.books.open(filename)
    sht = wb.sheets["sheet1"]

    rows = sht.used_range.shape[0]  # 获取列数

    # 读取五列的内容
    result = []
    for col in ["A", "B", "C", "D", "E"]:
        region = col + "2:" + col + str(rows)
        result.append(sht[region].value)

    wb.close()
    time.sleep(5)
    app.quit()
    return result
def read_info_pandas(filename):
    df = pd.read_excel(filename)

    result = []
    for i in range(5):
        result.append(df[df.columns[i]])

    return result

def write_invitation(target, description, department, number, plan, filename):

    university = "华清大学"
    prof = "张小明"
    sender = target + description + "有限公司"

    title = f"关于邀请{university}师生赴{target}进行暑期社会实践的函"
    receiver = f"共青团{university}委员会："
    text = f"获悉{university}正在组织优秀学生利用暑假赴海外国家进行社会实践的活动，作为在{target}的{description}，我们诚挚的邀请以{university}{department}{prof}教授带队的师生代表团共计{number}人来{target}进行社会实践，{plan}。"

    locale.setlocale(locale.LC_CTYPE, "chinese")
    date = time.strftime("%Y年%m月%d日")

    # 设置格式
    doc = docx.Document()

    for i in range(5):
        doc.add_paragraph("")

    para_heading = doc.add_paragraph("")  # 返回标题段落对象
    para_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 设置为居中对齐
    para_heading.paragraph_format.space_before = Pt(0)  # 设置段前间距 0 磅
    para_heading.paragraph_format.space_after = Pt(0)  # 设置段后 0 磅
    para_heading.paragraph_format.line_spacing = 1  # 设置行间距为 1 倍

    para_heading.paragraph_format.left_indent = Inches(0)  # 设置左缩进 0 英寸
    para_heading.paragraph_format.right_indent = Inches(0)  # 设置右缩进 0 英寸

    run = para_heading.add_run(title)
    run.font.name = "黑体"  # 设置字体
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")  # 设置为字体，和上边的保持一致
    run.font.size = Pt(20)  # 设置文字的大小为20磅
    run.font.color.rgb = RGBColor(0, 0, 0)  # 设置颜色为黑色

    # 设置正文格式
    doc.styles["Normal"].font.size = Pt(14)
    doc.styles["Normal"].font.name = "仿宋"  # 设置当文字是西文时的字体
    doc.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋")  # 设置当文字是中文时的字体

    # 添加内容
    for i in range(2):
        doc.add_paragraph("")

    doc.add_paragraph(receiver)

    paragraph = doc.add_paragraph(text)
    paragraph.paragraph_format.first_line_indent = Cm(1)  # 设置左缩进 0 英寸

    paragraph = doc.add_paragraph(sender)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    paragraph = doc.add_paragraph(date)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.save(filename)


def mkdir(folder_name):
    existance = os.path.exists(folder_name)  # 判断文件/文件夹是否存在

    if not existance:  # 若文件夹不存在则创建
        os.mkdir(folder_name)


# 读取 excel
file_source = "Data_samples/社会实践报名表.xlsx"
countries, descriptions, departments, numbers, plans = read_info_pandas(file_source)


# 若不存在则创建目录
folder_name = "邀请函"
mkdir(folder_name)

# 写入 word
for i in range(len(countries)):
    target = countries[i]
    description = descriptions[i]
    department = departments[i]

    number = int(numbers[i])
    plan = plans[i]

    filename = f"./{folder_name}/社会实践邀请函_{target}.docx"

    print(filename)
    write_invitation(target, description, department, number, plan, filename)
